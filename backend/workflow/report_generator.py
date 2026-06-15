"""Workflow Report Generator — generates DOCX/PDF/ODF reports for debate sessions.

Produces a structured report including:
- Debate title, case description, metadata
- Per-round agent outputs (full transcript)
- Consensus progression
- Final consensus and result
- Audit trail table (optional, if entries exist)
"""

from __future__ import annotations

import ast
import asyncio
import html as html_mod
import json
import logging
from datetime import datetime
from pathlib import Path
from typing import Any

from docx import Document
from weasyprint import HTML

from backend.workflow.audit_logger import AuditLogger
from backend.workflow.state_snapshot import StateSnapshotStore

logger = logging.getLogger(__name__)

_REPORTS_DIR = Path("reports")


def _display_agent_role(role: str) -> str:
    """Format an agent role for display.

    MVP debates pass pre-formatted names like ``"Strategist (deepseek-v4-flash)"``
    while legacy debates pass raw role names like ``"critic"``.  Capitalize
    only when the role starts with a lowercase letter.
    """
    if not role:
        return "Unbekannt"
    return role if role[0].isupper() else role.capitalize()


def _format_content_for_display(content: str, role: str = "") -> str:
    """Format agent content for human-readable display in exports.

    Orchestrator/meta-agent nodes often output structured JSON with fields
    like ``reasoning``, ``debate_status``, ``next_agent``, etc.  This helper
    detects such JSON and renders it as readable text.  Non-JSON content is
    returned as-is.
    """
    if not content:
        return ""

    stripped = content.strip()
    if not stripped.startswith("{"):
        return content

    try:
        data = json.loads(stripped)
    except (json.JSONDecodeError, TypeError):
        return content

    if not isinstance(data, dict):
        return content

    # Detect orchestrator-style JSON (has reasoning + some workflow fields)
    _orchestrator_keys = {"reasoning", "debate_status", "phase_transition", "next_agent", "contextual_directive", "injection_context"}
    if not _orchestrator_keys & set(data.keys()):
        return content

    # Render as readable text
    parts: list[str] = []
    label_map = {
        "reasoning": "Reasoning",
        "debate_status": "Status",
        "phase_transition": "Phase Transition",
        "next_agent": "Next Agent",
        "contextual_directive": "Directive",
        "injection_context": "Context",
    }
    for key in ("reasoning", "contextual_directive", "injection_context", "debate_status", "phase_transition", "next_agent"):
        value = data.get(key)
        if value is not None and str(value).strip():
            label = label_map.get(key, key.replace("_", " ").title())
            parts.append(f"**{label}:** {value}")

    return "\n\n".join(parts) if parts else content


def _build_node_phase_map(state: dict[str, Any], workflow_id: str = "") -> dict[str, dict[str, str]]:
    """Build node_id → {phase_name, phase_index} from workflow definition + node_sequence.

    Walks the ``node_sequence`` from the state snapshot.  When a ``wf-phase``
    node is encountered, the current phase context is updated.  All subsequent
    agent/gate nodes are assigned to that phase until the next phase node.

    Returns a dict keyed by node_id with ``phase_name`` (str) and
    ``phase_index`` (int, 1-based).
    """
    node_sequence: list[str] = state.get("node_sequence", [])
    if not node_sequence:
        return {}

    # Resolve phase_configs from the workflow definition
    phase_names: dict[str, str] = {}
    phase_set: set[str] = set()
    try:
        from backend.blueprints.repository import BlueprintRepository

        repo = BlueprintRepository()
        wf_def = repo.get_workflow_definition(workflow_id) if workflow_id else None
        if wf_def and wf_def.phase_configs:
            for pc in wf_def.phase_configs.values():
                phase_names[pc.phase_node_id] = pc.name
                phase_set.add(pc.phase_node_id)
    except Exception:
        pass

    if not phase_set:
        return {}

    result: dict[str, dict[str, str]] = {}
    current_phase = ""
    phase_counter = 0
    for nid in node_sequence:
        if nid in phase_set:
            current_phase = phase_names.get(nid, f"Phase {phase_counter + 1}")
            phase_counter += 1
        elif current_phase:
            result[nid] = {"phase_name": current_phase, "phase_index": phase_counter}
    return result


def _build_mvp_rounds_from_snapshot(debate_data: dict[str, Any]) -> list[dict[str, Any]]:
    """Build round data from state snapshot for MVP debates.

    MVP debates don't populate ``debate_data["rounds"]`` — the agent
    outputs live exclusively in the state snapshot's ``node_outputs``.
    This helper loads the snapshot, resolves agent names and LLM
    profiles from ``node_configs``, and returns a list of round dicts
    compatible with the existing renderer expectations.

    When the workflow has ``phase_configs``, the round dicts include a
    ``phases`` key with agent outputs grouped by debate phase, enabling
    hierarchical rendering in exports.

    Returns:
        A list of round dicts with keys ``round``, ``consensus``,
        ``agent_outputs``, and optionally ``phases``.
    """
    session_id = debate_data.get("session_id", "")
    if not session_id:
        return []

    try:
        snap_store = StateSnapshotStore()
        snapshot = snap_store.get_latest(session_id)
    except Exception:
        logger.warning("Could not load state snapshot for session %s", session_id)
        return []

    if not snapshot:
        return []

    state = snapshot.get("state", {})
    node_outputs: list[dict] = state.get("node_outputs", [])
    node_configs_raw: dict[str, Any] = state.get("node_configs", {})
    llm_assignments: dict[str, str] = debate_data.get("llm_assignments", {})
    workflow_id = state.get("workflow_id", "")

    if not node_outputs:
        return []

    # Parse node_configs — values may be stored as Python repr strings
    config_by_node: dict[str, dict] = {}
    for nid, cfg in node_configs_raw.items():
        if isinstance(cfg, str):
            try:
                cfg = ast.literal_eval(cfg)
            except (ValueError, SyntaxError):
                cfg = {}
        config_by_node[nid] = cfg if isinstance(cfg, dict) else {}

    # Build a UUID → profile name cache for fallback resolution
    _profile_name_cache: dict[str, str] = {}

    def _resolve_profile_name(profile_id: str) -> str:
        """Resolve an LLM profile UUID to a human-readable name."""
        if not profile_id:
            return ""
        if profile_id in _profile_name_cache:
            return _profile_name_cache[profile_id]
        try:
            from backend.services.profile_service import ProfileService

            ps = ProfileService()
            profile = ps.get_llm_profile(profile_id)
            if profile and profile.name:
                _profile_name_cache[profile_id] = profile.name
                return profile.name
        except Exception:
            pass
        _profile_name_cache[profile_id] = ""
        return ""

    # Build node_id → phase mapping for phase-aware workflows
    node_phase_map = _build_node_phase_map(state, workflow_id)

    # Build agent outputs with proper names and metadata
    agent_outputs: list[dict[str, Any]] = []
    for no in node_outputs:
        role = no.get("role", "")
        node_id = no.get("node_id", "")
        config = config_by_node.get(node_id, {})

        llm_model = config.get("llm_model", "")
        llm_profile_id = config.get("llm_profile_id", "") or llm_assignments.get(role, "")
        llm_profile_name = config.get("llm_profile_name", "")
        # Fallback: resolve UUID → name if llm_profile_name is missing
        if not llm_profile_name and llm_profile_id:
            llm_profile_name = _resolve_profile_name(llm_profile_id)
        role_type_name = config.get("role_type_name", "")

        # Build a human-readable agent name: "Strategist (deepseek-v4-flash)"
        agent_name = role_type_name or role.replace("_", " ").title()
        if llm_model:
            agent_name = f"{agent_name} ({llm_model})"
        elif llm_profile_name:
            agent_name = f"{agent_name} ({llm_profile_name})"

        phase_info = node_phase_map.get(node_id, {})
        agent_outputs.append(
            {
                "role": agent_name,
                "content": _format_content_for_display(no.get("content", ""), role),
                "tokens_used": no.get("tokens_used", 0),
                "duration_ms": no.get("duration_ms", 0),
                "llm_profile_id": llm_profile_id,
                "llm_profile_name": llm_profile_name,
                "round": no.get("round"),
                "phase_name": phase_info.get("phase_name", ""),
                "phase_index": phase_info.get("phase_index", 0),
            }
        )

    # Group by round number
    from collections import OrderedDict

    rounds_map: OrderedDict[int, list[dict[str, Any]]] = OrderedDict()
    for ao in agent_outputs:
        rnd = ao.get("round") or 0
        rounds_map.setdefault(rnd, []).append(ao)

    final_consensus = state.get("final_consensus", debate_data.get("result", {}).get("consensus", 0.0))
    current_round = state.get("current_round", debate_data.get("current_round", 1))

    # If no round info was available, fall back to single round
    if not rounds_map:
        return [
            {
                "round": current_round,
                "consensus": final_consensus,
                "agent_outputs": agent_outputs,
            }
        ]

    result: list[dict[str, Any]] = []
    has_phases = bool(node_phase_map)

    for rnd_num, rnd_outputs in rounds_map.items():
        round_dict: dict[str, Any] = {
            "round": rnd_num if rnd_num else current_round,
            "consensus": final_consensus,
            "agent_outputs": rnd_outputs,
        }

        # Group by phase within this round
        if has_phases:
            phase_groups: OrderedDict[str, dict[str, Any]] = OrderedDict()
            for ao in rnd_outputs:
                pname = ao.get("phase_name", "")
                if not pname:
                    continue
                if pname not in phase_groups:
                    phase_groups[pname] = {
                        "phase_name": pname,
                        "phase_index": ao.get("phase_index", 0),
                        "agent_outputs": [],
                    }
                phase_groups[pname]["agent_outputs"].append(ao)

            if phase_groups:
                round_dict["phases"] = list(phase_groups.values())

        result.append(round_dict)

    return result


def _build_node_llm_name_map(session_id: str) -> dict[str, str]:
    """Build a node_id → llm_profile_name mapping from the state snapshot.

    Used to resolve UUIDs in audit log entries to human-readable names.
    Falls back to resolving ``llm_profile_id`` UUIDs via the ProfileService
    when ``llm_profile_name`` is not stored in node_configs (legacy snapshots).
    """
    try:
        snap_store = StateSnapshotStore()
        snapshot = snap_store.get_latest(session_id)
    except Exception:
        return {}

    if not snapshot:
        return {}

    state = snapshot.get("state", {})
    node_configs_raw = state.get("node_configs", {})
    result: dict[str, str] = {}
    # Collect UUIDs that need fallback resolution
    unresolved: dict[str, str] = {}  # node_id → llm_profile_id UUID
    for nid, cfg in node_configs_raw.items():
        if isinstance(cfg, str):
            try:
                cfg = ast.literal_eval(cfg)
            except (ValueError, SyntaxError):
                cfg = {}
        if isinstance(cfg, dict):
            name = cfg.get("llm_profile_name", "")
            if name:
                result[nid] = name
            else:
                profile_id = cfg.get("llm_profile_id", "")
                if profile_id:
                    unresolved[nid] = profile_id

    # Fallback: resolve UUIDs via ProfileService
    if unresolved:
        try:
            from backend.services.profile_service import ProfileService

            ps = ProfileService()
            for nid, profile_id in unresolved.items():
                profile = ps.get_llm_profile(profile_id)
                if profile and profile.name:
                    result[nid] = profile.name
        except Exception:
            logger.debug("Failed to resolve LLM profile UUIDs via ProfileService", exc_info=True)

    return result


def _enrich_audit_entries(
    audit_entries: list[dict[str, Any]],
    node_llm_names: dict[str, str],
) -> list[dict[str, Any]]:
    """Enrich audit entries with resolved LLM profile names."""
    for entry in audit_entries:
        node_id = entry.get("node_id", "")
        resolved = node_llm_names.get(node_id, "")
        if resolved:
            entry["llm_profile_name"] = resolved
    return audit_entries


def _build_audit_context_map(session_id: str) -> dict[str, dict[str, Any]]:
    """Build node_id → {round, phase, role_type_name} from state snapshot + workflow def.

    Returns a dict keyed by node_id.  Each value contains:
      - ``round`` (int | None): the round number the node ran in
      - ``phase`` (str): human-readable phase name (e.g. "Phase 1 — Strategists")
      - ``role_type_name`` (str): agent role display name
    """
    try:
        snap_store = StateSnapshotStore()
        snapshot = snap_store.get_latest(session_id)
    except Exception:
        return {}

    if not snapshot:
        return {}

    state = snapshot.get("state", {})
    node_outputs: list[dict] = state.get("node_outputs", [])
    node_configs_raw: dict[str, Any] = state.get("node_configs", {})
    node_sequence: list[str] = state.get("node_sequence", [])
    workflow_id = state.get("workflow_id", "")

    # --- Round: from node_outputs ---
    node_round_map: dict[str, int] = {}
    for no in node_outputs:
        nid = no.get("node_id", "")
        rnd = no.get("round", 0)
        if nid and rnd:
            node_round_map[nid] = rnd

    # --- Role type name: from node_configs ---
    node_role_map: dict[str, str] = {}
    for nid, cfg in node_configs_raw.items():
        if isinstance(cfg, str):
            try:
                cfg = ast.literal_eval(cfg)
            except (ValueError, SyntaxError):
                cfg = {}
        if isinstance(cfg, dict):
            role_name = cfg.get("role_type_name", "")
            if role_name:
                node_role_map[nid] = role_name

    # --- Phase: from workflow definition phase_configs + node_sequence ---
    node_phase_map: dict[str, str] = {}
    try:
        from backend.blueprints.repository import BlueprintRepository

        repo = BlueprintRepository()
        wf_def = repo.get_workflow_definition(workflow_id) if workflow_id else None
    except Exception:
        wf_def = None

    phase_names: dict[str, str] = {}  # phase_node_id → display name
    phase_set: set[str] = set()
    if wf_def and wf_def.phase_configs:
        for pc in wf_def.phase_configs.values():
            phase_names[pc.phase_node_id] = pc.name
            phase_set.add(pc.phase_node_id)

    # Walk the node_sequence to assign phases.
    # When we encounter a wf-phase node, update the current phase name.
    # All subsequent nodes belong to that phase until the next phase node.
    current_phase = ""
    phase_counter = 0
    for nid in node_sequence:
        if nid in phase_set:
            current_phase = phase_names.get(nid, f"Phase {phase_counter + 1}")
            phase_counter += 1
        elif current_phase:
            node_phase_map[nid] = current_phase

    # --- Merge into context map ---
    result: dict[str, dict[str, Any]] = {}
    all_node_ids = set(node_round_map) | set(node_role_map) | set(node_phase_map) | set(node_sequence)
    for nid in all_node_ids:
        result[nid] = {
            "round": node_round_map.get(nid),
            "phase": node_phase_map.get(nid, ""),
            "role_type_name": node_role_map.get(nid, ""),
        }
    return result


def _format_audit_content(output_content: str, event_type: str = "") -> str:
    """Parse raw audit output_content JSON and extract human-readable text.

    The ``output_content`` from the audit log stores the full state-update
    dict returned by the node function.  This helper extracts only the
    meaningful text for display.
    """
    if not output_content:
        return ""

    # Try to parse as JSON
    data: Any = None
    try:
        data = json.loads(output_content)
    except (json.JSONDecodeError, TypeError):
        # Not JSON — return as-is (could be plain text from a failed node)
        return output_content.strip()

    if not isinstance(data, dict):
        return str(data).strip()

    # --- node_completed: extract agent output content ---
    # Agent nodes return {"current_draft": ..., "node_outputs": [...], ...}
    node_outputs = data.get("node_outputs")
    if isinstance(node_outputs, list) and node_outputs:
        first = node_outputs[0]
        if isinstance(first, dict):
            content = first.get("content", "")
            if content:
                return _format_content_for_display(content.strip())

    # --- Gate nodes: extract decision ---
    gate_decision = data.get("gate_decision")
    if isinstance(gate_decision, dict):
        decision = gate_decision.get("decision", "")
        reason = gate_decision.get("reason", "")
        parts = []
        if decision:
            parts.append(f"Entscheidung: {decision}")
        if reason:
            parts.append(f"Begründung: {reason}")
        if parts:
            return " | ".join(parts)

    # --- Input node: extract context ---
    context = data.get("context")
    if isinstance(context, str) and context:
        return context.strip()

    # --- Complete node: extract output ---
    output = data.get("output")
    if isinstance(output, str) and output:
        return output.strip()

    # --- Fallback: try any "content" key ---
    content = data.get("content")
    if isinstance(content, str) and content:
        return content.strip()

    # --- Last resort: compact JSON (truncated) ---
    try:
        compact = json.dumps(data, ensure_ascii=False, default=str)
        if len(compact) > 500:
            return compact[:500] + "…"
        return compact
    except Exception:
        return str(data)[:500]


class WorkflowReportGenerator:
    """Generates structured reports for completed debate sessions."""

    def __init__(self, db_path: Path | str | None = None) -> None:
        """Initialise WorkflowReportGenerator."""
        self._db_path = Path(db_path) if db_path else None
        self._audit = AuditLogger(self._db_path)
        _REPORTS_DIR.mkdir(parents=True, exist_ok=True)

    async def generate(
        self,
        session_id: str,
        fmt: str = "docx",
        debate_data: dict[str, Any] | None = None,
    ) -> Path:
        """Generate a report for *session_id* in the given *fmt*.

        Args:
            session_id: The workflow session ID (same as debate_id).
            fmt: Output format — ``"docx"``, ``"pdf"``, or ``"odf"``.
            debate_data: Optional debate dict from DebateStore.  When
                provided the report includes the full debate transcript.

        Returns:
            Path to the generated report file.
        """
        if fmt not in ("docx", "pdf", "odf"):
            raise ValueError(f"Unsupported format: {fmt!r}")

        ts = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"report_{session_id[:8]}_{ts}.{fmt}"
        path = _REPORTS_DIR / filename

        # Load audit data (supplementary)
        audit_entries = self._audit.get_audit_log_for_replay(session_id)

        # Build transcript data from debate store
        transcript = self._build_transcript(debate_data)

        # Resolve UUID → human-readable name for audit entries
        node_llm_names = _build_node_llm_name_map(session_id)
        _enrich_audit_entries(audit_entries, node_llm_names)

        if fmt == "docx":
            await asyncio.to_thread(self._build_docx, session_id, transcript, audit_entries, path)
        elif fmt == "pdf":
            await asyncio.to_thread(self._build_pdf, session_id, transcript, audit_entries, path)
        elif fmt == "odf":
            await asyncio.to_thread(self._build_odf, session_id, transcript, audit_entries, path)

        logger.info("Report generated: %s", path)
        return path

    # ------------------------------------------------------------------
    # Transcript extraction
    # ------------------------------------------------------------------

    @staticmethod
    def _build_transcript(debate_data: dict[str, Any] | None) -> dict[str, Any]:
        """Extract structured transcript data from a debate dict.

        Returns a normalized dict with keys: title, case_text, language,
        model, max_rounds, threshold, status, current_round,
        final_consensus, rounds (list of round dicts).
        """
        if not debate_data:
            return {
                "title": "",
                "case_text": "",
                "language": "de",
                "model": "",
                "max_rounds": 0,
                "threshold": 0.0,
                "status": "unknown",
                "current_round": 0,
                "final_consensus": 0.0,
                "rounds": [],
                "output": "",
            }

        req = debate_data.get("request", {})
        result = debate_data.get("result", {})
        rounds = debate_data.get("rounds", [])

        # rounds may come from either the debate store or the result dict
        if not rounds and result:
            rounds = result.get("rounds", [])

        # MVP debates: build rounds from state snapshot node_outputs
        is_mvp = debate_data.get("is_mvp", False)
        if is_mvp and not rounds:
            rounds = _build_mvp_rounds_from_snapshot(debate_data)

        return {
            "title": debate_data.get("title", ""),
            "case_text": req.get("case", {}).get("text", "") if isinstance(req.get("case"), dict) else str(req.get("case", "")),
            "language": req.get("language", "de"),
            "model": req.get("llm_profile_id", ""),
            "max_rounds": req.get("max_rounds", 0),
            "threshold": req.get("consensus_threshold", 0.0),
            "status": str(debate_data.get("status", "unknown")),
            "current_round": debate_data.get("current_round", result.get("current_round", 0)),
            "final_consensus": result.get("final_consensus", debate_data.get("final_consensus", 0.0)),
            "rounds": rounds,
            "output": result.get("output", ""),
            "is_mvp": is_mvp,
        }

    # ------------------------------------------------------------------
    # Transactional Drafting helpers
    # ------------------------------------------------------------------

    @staticmethod
    def _extract_transactional_protocol(
        audit_entries: list[dict[str, Any]],
    ) -> list[dict[str, Any]]:
        """Extract CriticItem → BuildResponse → PragmatistEvaluation data from audit entries.

        Parses ``builder_iteration`` and ``pragmatist_evaluation`` workflow
        events from the audit log and merges them into a unified protocol
        table for the report.
        """
        import json as _json

        iterations: dict[int, dict] = {}  # draft_version → builder data
        evaluations: list[dict] = []

        for entry in audit_entries:
            event_type = entry.get("event_type", "")
            raw_output = entry.get("output_content", "")

            if event_type == "builder_iteration":
                try:
                    if isinstance(raw_output, str) and raw_output:
                        data = _json.loads(raw_output)
                    elif isinstance(raw_output, dict):
                        data = raw_output
                    else:
                        data = {}
                except (ValueError, TypeError):
                    data = {}
                dv = entry.get("draft_version", data.get("draft_version", 0))
                iterations[dv] = {
                    "draft_version": dv,
                    "constructivity_score": entry.get("constructivity_score", data.get("constructivity_score")),
                    "build_response_count": data.get("build_response_count", 0),
                }

            elif event_type == "pragmatist_evaluation":
                try:
                    if isinstance(raw_output, str) and raw_output:
                        data = _json.loads(raw_output)
                    elif isinstance(raw_output, dict):
                        data = raw_output
                    else:
                        data = {}
                except (ValueError, TypeError):
                    data = {}
                verdicts = data.get("verdicts", [])
                for v in verdicts:
                    evaluations.append(v)

        # Build merged protocol entries
        protocol: list[dict[str, Any]] = []

        for ev in evaluations:
            resp_to = ev.get("response_to", "")
            verdict = ev.get("verdict", "pending")
            feasibility = ev.get("feasibility")

            # Find the iteration this belongs to
            draft_ver = 0
            for dv, it_data in sorted(iterations.items()):
                draft_ver = dv

            protocol.append(
                {
                    "critic_item_id": resp_to,
                    "build_response_summary": f"Verdict: {verdict}",
                    "verdict": verdict,
                    "feasibility": feasibility,
                    "draft_version": draft_ver,
                }
            )

        # If no pragmatist evaluations, fall back to builder iterations
        if not protocol and iterations:
            for dv, it_data in sorted(iterations.items()):
                protocol.append(
                    {
                        "critic_item_id": f"Iteration {dv}",
                        "build_response_summary": f"{it_data.get('build_response_count', 0)} Antworten",
                        "verdict": "pending",
                        "feasibility": it_data.get("constructivity_score"),
                        "draft_version": dv,
                    }
                )

        return protocol

    # ------------------------------------------------------------------
    # DOCX
    # ------------------------------------------------------------------

    @staticmethod
    def _docx_add_agent_output(doc: Document, ao: dict[str, Any], level: int) -> None:
        """Render a single agent output block into a DOCX document."""
        role = ao.get("role", "unbekannt")
        content = ao.get("content", "")
        tokens = ao.get("tokens_used", 0)
        duration_ms = ao.get("duration_ms", 0)
        llm_pid = ao.get("llm_profile_name", "") or ao.get("llm_profile_id", "")
        heading = _display_agent_role(role)
        if llm_pid:
            heading += f" — {llm_pid}"
        doc.add_heading(heading, level=level)
        if content:
            for paragraph in content.split("\n\n"):
                paragraph = paragraph.strip()
                if paragraph:
                    doc.add_paragraph(paragraph)
        meta_parts = []
        if tokens:
            meta_parts.append(f"Tokens: {tokens}")
        if duration_ms:
            meta_parts.append(f"Latenz: {duration_ms}ms")
        if llm_pid:
            meta_parts.append(f"LLM-Profil: {llm_pid}")
        if meta_parts:
            doc.add_paragraph(" | ".join(meta_parts)).italic = True

    def _build_docx(
        self,
        session_id: str,
        transcript: dict[str, Any],
        audit_entries: list[dict[str, Any]],
        path: Path,
    ) -> None:
        """Build docx internally."""
        doc = Document()
        doc.styles["Normal"].font.name = "Calibri"

        # --- Title ---
        title = transcript["title"] or f"Debatte {session_id[:8]}"
        doc.add_heading(title, level=0)

        # --- Metadata ---
        doc.add_heading("Zusammenfassung", level=1)
        table_meta = doc.add_table(rows=0, cols=2, style="Table Grid")
        meta_rows = [
            ("Session-ID", session_id),
            ("Titel", transcript["title"]),
            ("Sprache", transcript["language"]),
            ("Modell", transcript["model"]),
            ("Max. Runden", str(transcript["max_rounds"])),
            ("Konsens-Schwelle", f"{transcript['threshold']:.0%}"),
            ("Status", transcript["status"]),
            ("Runden absolviert", str(transcript["current_round"])),
            ("Finaler Konsens", f"{transcript['final_consensus']:.1%}"),
            ("Generiert", datetime.now().strftime("%Y-%m-%d %H:%M:%S")),
        ]
        for label, value in meta_rows:
            row = table_meta.add_row()
            row.cells[0].text = label
            row.cells[1].text = value

        # --- Case description ---
        case_text = transcript["case_text"]
        if case_text:
            doc.add_heading("Fallbeschreibung", level=1)
            doc.add_paragraph(case_text)

        # --- Debate transcript (per round) ---
        rounds = transcript["rounds"]
        if rounds:
            doc.add_heading("Debatte-Transkript", level=1)
            for rd in rounds:
                round_num = rd.get("round", "?")
                consensus = rd.get("consensus", 0.0)
                doc.add_heading(f"Runde {round_num} — Konsens: {consensus:.1%}", level=2)

                phases = rd.get("phases", [])
                if phases:
                    for phase in phases:
                        phase_name = phase.get("phase_name", "")
                        doc.add_heading(phase_name, level=3)
                        for ao in phase.get("agent_outputs", []):
                            self._docx_add_agent_output(doc, ao, level=4)
                else:
                    for ao in rd.get("agent_outputs", []):
                        self._docx_add_agent_output(doc, ao, level=3)

        # --- Final output ---
        output = transcript["output"]
        if output:
            doc.add_heading("Ergebnis", level=1)
            for paragraph in output.split("\n\n"):
                paragraph = paragraph.strip()
                if paragraph:
                    doc.add_paragraph(paragraph)

        # --- Transactional Drafting: Konstruktionsprotokoll ---
        txn_entries = self._extract_transactional_protocol(audit_entries)
        if txn_entries:
            doc.add_heading("Konstruktionsprotokoll", level=1)

            # Table: CriticItem → BuildResponse → PragmatistEvaluation
            table = doc.add_table(rows=1, cols=3, style="Table Grid")
            headers = ["CriticItem", "BuildResponse (Optionen)", "Pragmatist-Evaluation"]
            for i, h in enumerate(headers):
                table.rows[0].cells[i].text = h

            for entry in txn_entries:
                row = table.add_row()
                row.cells[0].text = entry.get("critic_item_id", "")
                row.cells[1].text = entry.get("build_response_summary", "")
                verdict = entry.get("verdict", "pending")
                feasibility = entry.get("feasibility")
                cell_text = verdict.upper()
                if feasibility is not None:
                    cell_text += f"\nMachbarkeit: {feasibility * 100:.1f}%"
                row.cells[2].text = cell_text

        # --- Audit trail (supplementary) ---
        if audit_entries:
            doc.add_heading("Audit-Trail", level=1)
            table = doc.add_table(rows=1, cols=7, style="Table Grid")
            headers = ["Zeitstempel", "Ereignis", "Knoten", "Aktor", "LLM-Profil", "Latenz (ms)", "Tokens"]
            for i, h in enumerate(headers):
                table.rows[0].cells[i].text = h
            for entry in audit_entries:
                row = table.add_row()
                row.cells[0].text = str(entry.get("timestamp", ""))
                row.cells[1].text = str(entry.get("event_type", ""))
                row.cells[2].text = str(entry.get("node_id", ""))
                row.cells[3].text = str(entry.get("actor", ""))
                row.cells[4].text = str(entry.get("llm_profile_name", "") or entry.get("llm_profile_id", ""))
                row.cells[5].text = str(entry.get("latency_ms", 0))
                row.cells[6].text = str(entry.get("prompt_tokens", 0) + entry.get("completion_tokens", 0))

        doc.save(str(path))

    # ------------------------------------------------------------------
    # PDF (via WeasyPrint)
    # ------------------------------------------------------------------

    def _build_pdf(
        self,
        session_id: str,
        transcript: dict[str, Any],
        audit_entries: list[dict[str, Any]],
        path: Path,
    ) -> None:
        """Build pdf internally."""
        html_content = self._render_html(session_id, transcript, audit_entries)
        HTML(string=html_content).write_pdf(str(path))

    # ------------------------------------------------------------------
    # ODF
    # ------------------------------------------------------------------

    def _build_odf(
        self,
        session_id: str,
        transcript: dict[str, Any],
        audit_entries: list[dict[str, Any]],
        path: Path,
    ) -> None:
        """Build odf internally."""
        html_content = self._render_html(session_id, transcript, audit_entries)
        try:
            from odf.opendocument import OpenDocumentText
            from odf.text import H, P

            doc = OpenDocumentText()

            title = transcript["title"] or f"Debatte {session_id[:8]}"
            doc.text.addElement(H(text=title, outlinelevel=1))

            # Metadata
            doc.text.addElement(H(text="Zusammenfassung", outlinelevel=2))
            meta_lines = [
                f"Session-ID: {session_id}",
                f"Titel: {transcript['title']}",
                f"Sprache: {transcript['language']}",
                f"Modell: {transcript['model']}",
                f"Max. Runden: {transcript['max_rounds']}",
                f"Konsens-Schwelle: {transcript['threshold']:.0%}",
                f"Status: {transcript['status']}",
                f"Runden absolviert: {transcript['current_round']}",
                f"Finaler Konsens: {transcript['final_consensus']:.1%}",
            ]
            for line in meta_lines:
                doc.text.addElement(P(text=line))

            # Case description
            case_text = transcript["case_text"]
            if case_text:
                doc.text.addElement(H(text="Fallbeschreibung", outlinelevel=2))
                for para in case_text.split("\n\n"):
                    para = para.strip()
                    if para:
                        doc.text.addElement(P(text=para))

            # Transcript
            rounds = transcript["rounds"]
            if rounds:
                doc.text.addElement(H(text="Debatte-Transkript", outlinelevel=2))
                for rd in rounds:
                    round_num = rd.get("round", "?")
                    consensus = rd.get("consensus", 0.0)
                    doc.text.addElement(H(text=f"Runde {round_num} — Konsens: {consensus:.1%}", outlinelevel=3))
                    phases = rd.get("phases", [])
                    if phases:
                        for phase in phases:
                            phase_name = phase.get("phase_name", "")
                            doc.text.addElement(H(text=phase_name, outlinelevel=4))
                            for ao in phase.get("agent_outputs", []):
                                role = ao.get("role", "unbekannt")
                                content = ao.get("content", "")
                                llm_pid = ao.get("llm_profile_name", "") or ao.get("llm_profile_id", "")
                                role_label = _display_agent_role(role)
                                if llm_pid:
                                    role_label += f" — {llm_pid}"
                                doc.text.addElement(P(text=f"[{role_label}]"))
                                if content:
                                    for para in content.split("\n\n"):
                                        para = para.strip()
                                        if para:
                                            doc.text.addElement(P(text=para))
                    else:
                        for ao in rd.get("agent_outputs", []):
                            role = ao.get("role", "unbekannt")
                            content = ao.get("content", "")
                            llm_pid = ao.get("llm_profile_name", "") or ao.get("llm_profile_id", "")
                            role_label = _display_agent_role(role)
                            if llm_pid:
                                role_label += f" — {llm_pid}"
                            doc.text.addElement(P(text=f"[{role_label}]"))
                            if content:
                                for para in content.split("\n\n"):
                                    para = para.strip()
                                    if para:
                                        doc.text.addElement(P(text=para))

            # Final output
            output = transcript["output"]
            if output:
                doc.text.addElement(H(text="Ergebnis", outlinelevel=2))
                for para in output.split("\n\n"):
                    para = para.strip()
                    if para:
                        doc.text.addElement(P(text=para))

            # Audit trail
            if audit_entries:
                doc.text.addElement(H(text="Audit-Trail", outlinelevel=2))
                for entry in audit_entries:
                    llm_pid = entry.get("llm_profile_name", "") or entry.get("llm_profile_id", "")
                    line = f"{entry.get('timestamp', '')} | {entry.get('event_type', '')} | {entry.get('node_id', '')} | {entry.get('actor', '')}"
                    if llm_pid:
                        line += f" | LLM: {llm_pid}"
                    doc.text.addElement(P(text=line))

            doc.save(str(path))
        except ImportError:
            # Fallback: write HTML if odfpy not available
            path = path.with_suffix(".html")
            path.write_text(html_content, encoding="utf-8")
            logger.warning("odfpy not available; wrote HTML instead: %s", path)

    # ------------------------------------------------------------------
    # HTML rendering (shared by PDF and ODF)
    # ------------------------------------------------------------------

    @staticmethod
    def _render_html(
        session_id: str,
        transcript: dict[str, Any],
        audit_entries: list[dict[str, Any]],
    ) -> str:
        """Render html the instance."""
        esc = html_mod.escape
        title = esc(transcript["title"] or f"Debatte {session_id[:8]}")

        # --- Metadata table ---
        meta_pairs = [
            ("Session-ID", session_id),
            ("Titel", esc(transcript["title"])),
            ("Sprache", esc(transcript["language"])),
            ("Modell", esc(transcript["model"])),
            ("Max. Runden", str(transcript["max_rounds"])),
            ("Konsens-Schwelle", f"{transcript['threshold']:.0%}"),
            ("Status", esc(transcript["status"])),
            ("Runden absolviert", str(transcript["current_round"])),
            ("Finaler Konsens", f"{transcript['final_consensus']:.1%}"),
            ("Generiert", datetime.now().strftime("%Y-%m-%d %H:%M:%S")),
        ]
        meta_rows = "".join(f"<tr><td><strong>{k}</strong></td><td>{v}</td></tr>" for k, v in meta_pairs)

        # --- Case description ---
        case_text = transcript["case_text"]
        case_section = ""
        if case_text:
            case_section = f'<h2>Fallbeschreibung</h2><div class="case-text">{esc(case_text)}</div>'

        # --- Rounds ---
        rounds_html = ""
        rounds = transcript["rounds"]
        if rounds:
            rounds_html += "<h2>Debatte-Transkript</h2>"
            for rd in rounds:
                round_num = rd.get("round", "?")
                consensus = rd.get("consensus", 0.0)
                rounds_html += f"<h3>Runde {round_num} — Konsens: {consensus:.1%}</h3>"

                def _render_agent_block(ao: dict) -> str:
                    """Render a single agent output as an HTML block."""
                    _role = ao.get("role", "unbekannt")
                    _content = esc(ao.get("content", ""))
                    _tokens = ao.get("tokens_used", 0)
                    _duration = ao.get("duration_ms", 0)
                    _llm = ao.get("llm_profile_name", "") or ao.get("llm_profile_id", "")
                    _display = esc(_display_agent_role(_role))
                    if _llm:
                        _display += f' <span style="font-weight:normal;color:#666;">— {esc(_llm)}</span>'
                    _html = f'<div class="agent-block"><div class="agent-role">{_display}</div><div class="agent-content">{_content}</div>'
                    _meta = []
                    if _tokens:
                        _meta.append(f"Tokens: {_tokens}")
                    if _duration:
                        _meta.append(f"Latenz: {_duration}ms")
                    if _llm:
                        _meta.append(f"LLM-Profil: {esc(_llm)}")
                    if _meta:
                        _html += f'<div class="agent-meta">{" | ".join(_meta)}</div>'
                    _html += "</div>"
                    return _html

                phases = rd.get("phases", [])
                if phases:
                    for phase in phases:
                        phase_name = phase.get("phase_name", "")
                        rounds_html += f'<h4 class="phase-heading">{esc(phase_name)}</h4>'
                        for ao in phase.get("agent_outputs", []):
                            rounds_html += _render_agent_block(ao)
                else:
                    for ao in rd.get("agent_outputs", []):
                        rounds_html += _render_agent_block(ao)

        # --- Final output ---
        output = transcript["output"]
        output_section = ""
        if output:
            output_section = f'<h2>Ergebnis</h2><div class="output-text">{esc(output)}</div>'

        # --- Transactional Drafting: Konstruktionsprotokoll ---
        txn_protocol = WorkflowReportGenerator._extract_transactional_protocol(audit_entries)
        protocol_section = ""
        if txn_protocol:
            protocol_rows = ""
            for entry in txn_protocol:
                verdict = entry.get("verdict", "pending")
                feasibility = entry.get("feasibility")
                feas_str = f"{feasibility * 100:.1f}%" if feasibility is not None else "—"
                bg = "#e8f5e9" if verdict == "accept" else "#fff3e0" if verdict == "revise" else "#ffebee" if verdict == "reject" else "#f5f5f5"
                color = "#2e7d32" if verdict == "accept" else "#e65100" if verdict == "revise" else "#c62828" if verdict == "reject" else "#757575"
                protocol_rows += (
                    f"<tr>"
                    f"<td>{esc(str(entry.get('critic_item_id', '')))}</td>"
                    f"<td>{esc(str(entry.get('build_response_summary', '')))}</td>"
                    f'<td style="background:{bg};color:{color};font-weight:bold;">{verdict.upper()}<br><small>{esc(feas_str)}</small></td>'
                    f"</tr>"
                )
            protocol_section = f"""<h2>Konstruktionsprotokoll</h2>
<table>
<thead><tr><th>CriticItem</th><th>BuildResponse</th><th>Pragmatist-Evaluation</th></tr></thead>
<tbody>{protocol_rows}</tbody>
</table>"""

        # --- Audit trail ---
        audit_rows = ""
        for e in audit_entries:
            audit_rows += (
                f"<tr>"
                f"<td>{esc(str(e.get('timestamp', '')))}</td>"
                f"<td>{esc(str(e.get('event_type', '')))}</td>"
                f"<td>{esc(str(e.get('node_id', '')))}</td>"
                f"<td>{esc(str(e.get('actor', '')))}</td>"
                f"<td>{esc(str(e.get('llm_profile_name', '') or e.get('llm_profile_id', '')))}</td>"
                f"<td>{e.get('latency_ms', 0)}</td>"
                f"<td>{e.get('prompt_tokens', 0) + e.get('completion_tokens', 0)}</td>"
                f"</tr>"
            )
        audit_section = ""
        if audit_entries:
            audit_section = f"""<h2>Audit-Trail</h2>
<table>
<thead><tr><th>Zeitstempel</th><th>Ereignis</th><th>Knoten</th><th>Aktor</th><th>LLM-Profil</th><th>Latenz (ms)</th><th>Tokens</th></tr></thead>
<tbody>{audit_rows}</tbody>
</table>"""

        return f"""<!DOCTYPE html>
<html><head><meta charset="utf-8">
<title>Bericht — {title}</title>
<style>
  body {{ font-family: Calibri, 'Segoe UI', sans-serif; margin: 2em; color: #222; line-height: 1.6; }}
  h1 {{ color: #1a1a2e; border-bottom: 2px solid #16213e; padding-bottom: 0.3em; }}
  h2 {{ color: #16213e; margin-top: 1.5em; border-bottom: 1px solid #ddd; padding-bottom: 0.2em; }}
  h3 {{ color: #0f3460; margin-top: 1.2em; }}
  h4 {{ color: #1a1a2e; margin-top: 1em; margin-bottom: 0.3em; font-size: 1.05em; }}
  .phase-heading {{ background: #f0f0f8; padding: 0.4em 0.8em; border-left: 4px solid #6366f1; border-radius: 4px; }}
  table {{ border-collapse: collapse; width: 100%; margin: 0.5em 0; }}
  th, td {{ border: 1px solid #ccc; padding: 6px 10px; text-align: left; }}
  th {{ background: #f0f0f0; }}
  .case-text {{ background: #f8f9fa; padding: 1em; border-left: 4px solid #0f3460; margin: 0.5em 0; white-space: pre-wrap; }}
  .agent-block {{ margin: 1em 0; padding: 0.8em; border: 1px solid #e0e0e0; border-radius: 6px; background: #fafafa; }}
  .agent-role {{ font-weight: bold; color: #0f3460; margin-bottom: 0.4em; font-size: 1.05em; }}
  .agent-content {{ white-space: pre-wrap; }}
  .agent-meta {{ font-size: 0.85em; color: #888; margin-top: 0.4em; }}
  .output-text {{ background: #e8f5e9; padding: 1em; border-left: 4px solid #2e7d32; white-space: pre-wrap; }}
  .meta-table {{ width: auto; }}
  .meta-table td {{ padding: 4px 12px; }}
</style></head><body>
<h1>{title}</h1>
<table class="meta-table">{meta_rows}</table>
{case_section}
{rounds_html}
{output_section}
{protocol_section}
{audit_section}
</body></html>"""
